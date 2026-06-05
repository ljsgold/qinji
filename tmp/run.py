# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import os

SRC = r'E:\python大作业\大作业1'
DOCX = os.path.join(SRC, 'output', 'doc', '刘俊昇_249050123_《Python程序设计》期末大作业任务书_可提交版 - 副本.docx')
OUT = os.path.join(SRC, 'output', 'doc', '刘俊昇_249050123_《Python程序设计》期末大作业任务书_丰富版.docx')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Helpers
def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.bold = True; r.font.size = Pt(22)
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(18)

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.bold = True; r.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)

def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.bold = True; r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

def bold_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.bold = True; r.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

def code(doc, text):
    for line in text.strip().split(chr(10)):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = 'Consolas'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        r.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2; p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)

def bullet(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(12)
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.5

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

def set_cell(cell, text, fn='宋体', fs=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(text)
    r.font.name = fn; r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    r.font.size = fs; r.bold = bold

def shade(row, color='D9E2F3'):
    for c in row.cells:
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd ' + nsdecls('w') + ' w:fill="' + color + '"/>'))

def make_table(doc, headers, rows, hdr_color='2F5496'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        set_cell(hdr.cells[i], h, '黑体', Pt(10), True)
    shade(hdr, hdr_color)
    for c in hdr.cells:
        for p in c.paragraphs:
            for r in p.runs: r.font.color.rgb = RGBColor(255,255,255)
    for ri, rd in enumerate(rows):
        row = t.rows[ri+1]
        for ci, v in enumerate(rd):
            set_cell(row.cells[ci], str(v), '宋体', Pt(9))
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci>0 else WD_ALIGN_PARAGRAPH.CENTER
        if ri%2==1: shade(row, 'E8EDF5')
    doc.add_paragraph()

print('Functions defined OK')

# ===== MAIN =====
doc = Document(DOCX)
print(f'Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

# Find report title
cut = None
for i, p in enumerate(doc.paragraphs):
    if '大作业报告' in p.text:
        cut = i; break
if cut is None:
    print('ERROR: cannot find title'); sys.exit(1)
print(f'Report title at paragraph index {cut}')

# Remove content after title
to_del = [p._element for i, p in enumerate(doc.paragraphs) if i >= cut]
for e in to_del: e.getparent().remove(e)
for t in doc.tables: t._element.getparent().remove(t._element)

print('Old content removed, adding new content...')

# Save as new file
doc.save(OUT)
print(f'Base structure saved to: {OUT}')
print('Base structure OK, now run enrich_full.py for full content')
